#!/usr/bin/env python3
"""
Audio Transcription GUI - Modern Tkinter Application with ttkbootstrap

Features:
- Dark/Light theme toggle
- Drag & drop audio files
- Cancel transcription support
- Keyboard shortcuts & tooltips
- Window state persistence
- Export formats (TXT, SRT, VTT, DOCX, JSON)
- Settings dialog
- Batch search/filter & multi-select
- Transcription history panel
- Audio playback preview
- Progress cancellation with threading
- Accessibility improvements (focus indicators, tooltips)
"""

import json
import os
import queue
import threading
import time
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk
from typing import Any, Callable, Optional

import transcription as t

# Optional dependencies
try:
    import ttkbootstrap as tb
    from ttkbootstrap.constants import *
    from ttkbootstrap import Style
    try:
        from ttkbootstrap.widgets import ToolTip
    except ImportError:
        # Fallback for older ttkbootstrap versions
        from ttkbootstrap.tooltip import ToolTip
    TTKBOOTSTRAP_AVAILABLE = True
except ImportError:
    TTKBOOTSTRAP_AVAILABLE = False
    tb = ttk
    ToolTip = None
    Style = None

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    # Check if tkdnd library is available
    import tkinter as tk
    test_root = tk.Tk()
    test_root.withdraw()
    try:
        test_root.tk.call('package', 'require', 'tkdnd')
        DND_AVAILABLE = True
    except tk.TclError:
        DND_AVAILABLE = False
    test_root.destroy()
except ImportError:
    DND_AVAILABLE = False

try:
    import pygame
    PYGAME_AVAILABLE = True
except ImportError:
    PYGAME_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─── Configuration & Constants ────────────────────────────────────────────

SETTINGS_FILE = Path.home() / ".simple_transcription_settings.json"
DEFAULT_SETTINGS = {
    "theme": "darkly" if TTKBOOTSTRAP_AVAILABLE else "default",
    "window_geometry": "1000x700+100+100",
    "last_tab": 0,
    "language": "Auto (detect language)",
    "auto_save_speaker": True,
    "default_export_format": "txt",
    "audio_preview_enabled": PYGAME_AVAILABLE,
    "font_size": 10,
    "remember_window_state": True,
    "auto_refresh_batch": True,
    "confirm_before_cancel": True,
}

EXPORT_FORMATS = [
    ("Text (.txt)", "txt"),
    ("SubRip Subtitle (.srt)", "srt"),
    ("WebVTT (.vtt)", "vtt"),
    ("JSON (.json)", "json"),
] + ([("Word Document (.docx)", "docx")] if DOCX_AVAILABLE else [])


# ─── Utility Classes ──────────────────────────────────────────────────────

class Settings:
    """Persistent settings manager."""
    
    def __init__(self):
        self._data = DEFAULT_SETTINGS.copy()
        self.load()
    
    def load(self):
        if SETTINGS_FILE.exists():
            try:
                with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                    loaded = json.load(f)
                    self._data.update(loaded)
            except Exception:
                pass
    
    def save(self):
        try:
            with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
                json.dump(self._data, f, indent=2)
        except Exception:
            pass
    
    def get(self, key: str, default=None):
        return self._data.get(key, default)
    
    def set(self, key: str, value: Any):
        self._data[key] = value
        self.save()
    
    def __getitem__(self, key: str):
        return self._data[key]
    
    def __setitem__(self, key: str, value: Any):
        self.set(key, value)


class CancellationToken:
    """Thread-safe cancellation token for long-running operations."""
    
    def __init__(self):
        self._cancelled = False
        self._lock = threading.Lock()
    
    def cancel(self):
        with self._lock:
            self._cancelled = True
    
    def reset(self):
        with self._lock:
            self._cancelled = False
    
    @property
    def is_cancelled(self) -> bool:
        with self._lock:
            return self._cancelled


class AudioPlayer:
    """Simple audio playback using pygame."""
    
    def __init__(self):
        self._initialized = False
        self._current_file = None
    
    def init(self):
        if PYGAME_AVAILABLE and not self._initialized:
            try:
                pygame.mixer.init()
                self._initialized = True
            except Exception:
                self._initialized = False
    
    def play(self, filepath: str):
        if not self._initialized:
            self.init()
        if not self._initialized:
            return False
        try:
            if self._current_file != filepath:
                pygame.mixer.music.load(filepath)
                self._current_file = filepath
            pygame.mixer.music.play()
            return True
        except Exception:
            return False
    
    def pause(self):
        if self._initialized:
            try:
                pygame.mixer.music.pause()
            except Exception:
                pass
    
    def stop(self):
        if self._initialized:
            try:
                pygame.mixer.music.stop()
            except Exception:
                pass
    
    def is_playing(self) -> bool:
        if self._initialized:
            try:
                return pygame.mixer.music.get_busy()
            except Exception:
                pass
        return False
    
    def set_volume(self, volume: float):
        if self._initialized:
            try:
                pygame.mixer.music.set_volume(max(0.0, min(1.0, volume)))
            except Exception:
                pass


class ToolTipManager:
    """Manages tooltips for widgets (fallback if ttkbootstrap not available)."""
    
    def __init__(self, root: tk.Tk):
        self.root = root
        self.tooltips = {}
    
    def add(self, widget: tk.Widget, text: str):
        if TTKBOOTSTRAP_AVAILABLE and ToolTip:
            try:
                # Only use ttkbootstrap ToolTip if style has a valid active theme object
                style_inst = Style.get_instance() if Style else None
                if style_inst and style_inst.theme is not None:
                    ToolTip(widget, text=text, bootstyle="info")
                else:
                    self._add_native_tooltip(widget, text)
            except Exception:
                # Fall back to native tooltip if ttkbootstrap tooltip fails
                self._add_native_tooltip(widget, text)
        else:
            self._add_native_tooltip(widget, text)
    
    def _add_native_tooltip(self, widget: tk.Widget, text: str):
        tooltip = None
        
        def show(event):
            nonlocal tooltip
            if tooltip:
                return
            tooltip = tk.Toplevel(widget)
            tooltip.wm_overrideredirect(True)
            tooltip.wm_attributes("-topmost", True)
            label = tk.Label(tooltip, text=text, background="#ffffe0", 
                           relief="solid", borderwidth=1, font=("TkDefaultFont", 9))
            label.pack()
            x = widget.winfo_rootx() + 20
            y = widget.winfo_rooty() + widget.winfo_height() + 5
            tooltip.geometry(f"+{x}+{y}")
        
        def hide(event):
            nonlocal tooltip
            if tooltip:
                tooltip.destroy()
                tooltip = None
        
        widget.bind("<Enter>", show)
        widget.bind("<Leave>", hide)
        self.tooltips[widget] = (show, hide)


# ─── Main Application ─────────────────────────────────────────────────────

class TranscriptionApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.settings = Settings()
        
        # Core state
        self.event_queue: queue.Queue = queue.Queue()
        self.busy = False
        self.api_key_configured = False
        self.speaker_conversation: list[dict] | None = None
        self.transcription_start_time: float | None = None
        self.current_transcription_file: str | None = None
        self.cancellation_token = CancellationToken()
        self.batch_worker_thread: threading.Thread | None = None
        self.speaker_worker_thread: threading.Thread | None = None
        
        # UI components
        self.audio_player = AudioPlayer()
        self.tooltip_manager = ToolTipManager(root)
        self.batch_tree_items = {}  # item_id -> audio_name
        self.history_items = []  # List of (timestamp, filename, status, duration)
        
        # Apply settings
        self._apply_settings()
        
        # Build UI
        self._build_ui()
        
        # Initialize
        self._try_load_api_key_from_env()
        self._update_app_state()
        self._poll_queue()
        self._start_elapsed_timer()
        self._restore_window_state()
        
        # Bind events
        self._bind_shortcuts()
        self._setup_drag_drop()
        
        # Protocol handlers
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        self.root.bind("<Configure>", self._on_window_configure)
    
    # ─── Settings & Persistence ──────────────────────────────────────────
    
    def _apply_settings(self):
        """Apply persisted settings to the application."""
        theme = self.settings.get("theme", "darkly")
        if TTKBOOTSTRAP_AVAILABLE:
            try:
                # Ensure the theme is a valid ttkbootstrap theme to prevent crash
                if theme in self.root.style.theme_names():
                    self.root.style.theme_use(theme)
                else:
                    # Fallback if theme from settings is standard Tkinter theme (e.g. clam, default)
                    fallback_theme = "darkly"
                    self.root.style.theme_use(fallback_theme)
                    self.settings["theme"] = fallback_theme
            except Exception:
                pass
        
        geometry = self.settings.get("window_geometry")
        if geometry and self.settings.get("remember_window_state", True):
            try:
                self.root.geometry(geometry)
            except Exception:
                pass
        
        font_size = self.settings.get("font_size", 10)
        self.root.option_add("*Font", f"TkDefaultFont {font_size}")
    
    def _restore_window_state(self):
        """Restore tab selection and other UI state."""
        last_tab = self.settings.get("last_tab", 0)
        if hasattr(self, 'notebook') and self.notebook.index("end") > last_tab:
            self.notebook.select(last_tab)
        
        language = self.settings.get("language", t.AUTO_LANGUAGE_LABEL)
        if language in t.LANGUAGE_COMBO_VALUES:
            self.language_var.set(language)
    
    def _save_window_state(self):
        """Save current window state to settings."""
        if self.settings.get("remember_window_state", True):
            self.settings["window_geometry"] = self.root.geometry()
            if hasattr(self, 'notebook'):
                self.settings["last_tab"] = self.notebook.index(self.notebook.select())
    
    def _on_window_configure(self, event):
        if event.widget == self.root:
            self._save_window_state()
    
    def _on_close(self):
        self._save_window_state()
        if self.busy:
            if self.settings.get("confirm_before_cancel", True):
                if not messagebox.askyesno("Confirm Exit", 
                    "Transcription in progress. Cancel and exit?", parent=self.root):
                    return
            self.cancellation_token.cancel()
        self.root.destroy()
    
    # ─── UI Construction ─────────────────────────────────────────────────
    
    def _build_ui(self):
        """Build the complete UI."""
        # Main container with padding
        self.main_container = ttk.Frame(self.root, padding=10)
        self.main_container.pack(fill="both", expand=True)
        
        # API Key Panel (initially hidden)
        self._build_api_key_panel()
        
        # Top toolbar with theme, language, settings
        self._build_toolbar()
        
        # Notebook with tabs
        self._build_notebook()
        
        # Status bar at bottom
        self._build_status_bar()
    
    def _build_api_key_panel(self):
        """API key entry panel shown when key is missing."""
        self.api_key_frame = ttk.LabelFrame(
            self.main_container,
            text="⚠ API Key Required",
            padding=15,
        )
        
        icon_label = ttk.Label(self.api_key_frame, text="🔑", font=("TkDefaultFont", 24))
        icon_label.pack(pady=(0, 10))
        
        ttk.Label(
            self.api_key_frame,
            text=(
                "No AssemblyAI API key found.\n"
                "Paste your key below and click Save to continue."
            ),
            justify="center",
            wraplength=500,
        ).pack(pady=(0, 15))
        
        self.api_key_text = tk.Text(self.api_key_frame, height=2, width=60, wrap="word",
                                     font=("Consolas", 10))
        self.api_key_text.pack(pady=(0, 10))
        self.tooltip_manager.add(self.api_key_text, "Paste your AssemblyAI API key here")
        
        options = ttk.Frame(self.api_key_frame)
        options.pack(fill="x", pady=(0, 10))
        
        self.save_env_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            options,
            text="Save to .env file for future sessions",
            variable=self.save_env_var,
        ).pack(side="left")
        
        btn_frame = ttk.Frame(self.api_key_frame)
        btn_frame.pack(fill="x")
        
        ttk.Button(
            btn_frame,
            text="Get API Key →",
            command=lambda: self._open_url("https://www.assemblyai.com/dashboard/signup"),
            bootstyle="link" if TTKBOOTSTRAP_AVAILABLE else None,
        ).pack(side="left")
        
        ttk.Button(
            btn_frame,
            text="Save & Continue",
            command=self._submit_api_key,
            bootstyle="success" if TTKBOOTSTRAP_AVAILABLE else None,
        ).pack(side="right", padx=(10, 0))
    
    def _build_toolbar(self):
        """Top toolbar with theme toggle, language, settings."""
        self.toolbar = ttk.Frame(self.main_container)
        self.toolbar.pack(fill="x", pady=(0, 10))
        
        # Left side: Theme toggle
        left_frame = ttk.Frame(self.toolbar)
        left_frame.pack(side="left")
        
        ttk.Label(left_frame, text="🎨 Theme:").pack(side="left", padx=(0, 5))
        
        self.theme_var = tk.StringVar(value=self.settings.get("theme", "darkly"))
        themes = ["darkly", "superhero", "solar", "cyborg", "vapor", "flatly", "litera", "minty", "lumen", "sandstone", "yeti", "pulse", "morph", "journal"]
        if not TTKBOOTSTRAP_AVAILABLE:
            themes = ["default", "clam", "alt", "classic"]
        
        self.theme_combo = ttk.Combobox(
            left_frame,
            textvariable=self.theme_var,
            values=themes,
            state="readonly",
            width=14,
        )
        self.theme_combo.pack(side="left")
        self.theme_combo.bind("<<ComboboxSelected>>", self._on_theme_change)
        self.tooltip_manager.add(self.theme_combo, "Switch UI theme (requires restart for some themes)")
        
        # Center: Language selection
        center_frame = ttk.Frame(self.toolbar)
        center_frame.pack(side="left", expand=True, padx=20)
        
        ttk.Label(center_frame, text="🌐 Language:").pack(side="left", padx=(0, 5))
        
        self.language_var = tk.StringVar(value=t.AUTO_LANGUAGE_LABEL)
        self.language_combo = ttk.Combobox(
            center_frame,
            textvariable=self.language_var,
            values=t.LANGUAGE_COMBO_VALUES,
            state="readonly",
            width=30,
        )
        self.language_combo.pack(side="left")
        self.language_combo.bind("<<ComboboxSelected>>", lambda e: self._on_language_change())
        self.tooltip_manager.add(self.language_combo, "Select transcription language (Auto = detect)")
        
        # Right side: Settings button
        right_frame = ttk.Frame(self.toolbar)
        right_frame.pack(side="right")
        
        ttk.Button(
            right_frame,
            text="⚙ Settings",
            command=self._open_settings,
            bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None,
        ).pack(side="right", padx=(5, 0))
        self.tooltip_manager.add(right_frame.winfo_children()[-1], "Open settings dialog (Ctrl+,)")
    
    def _build_notebook(self):
        """Main notebook with Batch and Speaker tabs + History."""
        self.notebook = ttk.Notebook(self.main_container)
        self.notebook.pack(fill="both", expand=True)
        self.notebook.bind("<<NotebookTabChanged>>", self._on_tab_change)
        
        # Batch Transcription Tab
        self.batch_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(self.batch_frame, text="📁 Batch Transcription")
        self._build_batch_tab(self.batch_frame)
        
        # Speaker-labeled Tab
        self.speaker_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(self.speaker_frame, text="🎙 Speaker Labeled")
        self._build_speaker_tab(self.speaker_frame)
        
        # History Tab
        self.history_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(self.history_frame, text="📜 History")
        self._build_history_tab(self.history_frame)
    
    def _build_batch_tab(self, parent: ttk.Frame):
        """Batch transcription tab with file list, progress, log."""
        # Toolbar
        toolbar = ttk.Frame(parent)
        toolbar.pack(fill="x", pady=(0, 8))
        
        # Left group
        left_group = ttk.Frame(toolbar)
        left_group.pack(side="left")
        
        self.refresh_button = ttk.Button(
            left_group, text="🔄 Refresh", command=self.refresh_batch_list
        )
        self.refresh_button.pack(side="left")
        self.tooltip_manager.add(self.refresh_button, "Refresh file list (F5)")
        
        self.transcribe_batch_button = ttk.Button(
            left_group, text="▶ Start Transcription", command=self.start_batch_transcription,
            bootstyle="success" if TTKBOOTSTRAP_AVAILABLE else None
        )
        self.transcribe_batch_button.pack(side="left", padx=(8, 0))
        self.tooltip_manager.add(self.transcribe_batch_button, "Start batch transcription (Ctrl+Enter)")
        
        self.cancel_batch_button = ttk.Button(
            left_group, text="⏹ Cancel", command=self._cancel_batch_transcription,
            state="disabled", bootstyle="danger" if TTKBOOTSTRAP_AVAILABLE else None
        )
        self.cancel_batch_button.pack(side="left", padx=(8, 0))
        self.tooltip_manager.add(self.cancel_batch_button, "Cancel current transcription (Esc)")
        
        # Right group
        right_group = ttk.Frame(toolbar)
        right_group.pack(side="right")
        
        self.view_transcript_button = ttk.Button(
            right_group, text="👁 View Transcript", command=self.view_transcript, state="disabled"
        )
        self.view_transcript_button.pack(side="left", padx=(0, 8))
        self.tooltip_manager.add(self.view_transcript_button, "View selected transcript (Enter)")
        
        self.export_batch_button = ttk.Button(
            right_group, text="📤 Export", command=self._export_batch_transcripts, state="disabled"
        )
        self.export_batch_button.pack(side="left")
        self.tooltip_manager.add(self.export_batch_button, "Export selected transcripts (Ctrl+E)")
        
        # Search/Filter bar
        search_frame = ttk.Frame(parent)
        search_frame.pack(fill="x", pady=(0, 8))
        
        ttk.Label(search_frame, text="🔍 Filter:").pack(side="left")
        
        self.batch_search_var = tk.StringVar()
        self.batch_search_var.trace_add("write", lambda *_: self._filter_batch_list())
        search_entry = ttk.Entry(search_frame, textvariable=self.batch_search_var, width=30)
        search_entry.pack(side="left", padx=(5, 10), fill="x", expand=True)
        self.tooltip_manager.add(search_entry, "Filter file list (type to search)")
        
        self.batch_status_label = ttk.Label(search_frame, text="Ready", foreground="gray")
        self.batch_status_label.pack(side="right")
        
        # Pack tree and scrollbars frame
        tree_frame = ttk.Frame(parent)
        tree_frame.pack(fill="both", expand=True, pady=(0, 8))

        # File list with columns
        columns = ("file", "status", "size", "modified", "duration")
        self.batch_tree = ttk.Treeview(
            tree_frame, columns=columns, show="headings", height=12, selectmode="extended"
        )
        self.batch_tree.heading("file", text="📄 File", command=lambda: self._sort_tree("file"))
        self.batch_tree.heading("status", text="Status", command=lambda: self._sort_tree("status"))
        self.batch_tree.heading("size", text="Size", command=lambda: self._sort_tree("size"))
        self.batch_tree.heading("modified", text="Modified", command=lambda: self._sort_tree("modified"))
        self.batch_tree.heading("duration", text="Duration", command=lambda: self._sort_tree("duration"))
        
        self.batch_tree.column("file", width=300, anchor="w", stretch=True)
        self.batch_tree.column("status", width=120, anchor="center")
        self.batch_tree.column("size", width=80, anchor="center")
        self.batch_tree.column("modified", width=140, anchor="center")
        self.batch_tree.column("duration", width=80, anchor="center")
        
        # Scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.batch_tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.batch_tree.xview)
        self.batch_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        self.batch_tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Progress bar
        progress_frame = ttk.Frame(parent)
        progress_frame.pack(fill="x", pady=(0, 4))
        
        self.batch_progress = ttk.Progressbar(progress_frame, mode="determinate")
        self.batch_progress.pack(fill="x")
        
        # Progress status
        status_frame = ttk.Frame(progress_frame)
        status_frame.pack(fill="x", pady=(4, 0))
        
        self.batch_current_label = ttk.Label(status_frame, text="")
        self.batch_current_label.pack(side="left")
        
        self.batch_elapsed_label = ttk.Label(status_frame, text="", foreground="gray")
        self.batch_elapsed_label.pack(side="right")
        
        # Log area
        log_label_frame = ttk.Frame(parent)
        log_label_frame.pack(fill="x", pady=(8, 4))
        
        ttk.Label(log_label_frame, text="📋 Transcription Log").pack(side="left")
        
        log_btn_frame = ttk.Frame(log_label_frame)
        log_btn_frame.pack(side="right")
        
        ttk.Button(log_btn_frame, text="Clear", command=self._clear_batch_log,
                  bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="left", padx=2)
        ttk.Button(log_btn_frame, text="Copy All", command=self._copy_batch_log,
                  bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="left", padx=2)
        ttk.Button(log_btn_frame, text="Save Log", command=self._save_batch_log,
                  bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="left", padx=2)
        
        self.batch_log = scrolledtext.ScrolledText(parent, height=8, state="disabled", wrap="word",
                                                    font=("Consolas", 9))
        self.batch_log.pack(fill="both", expand=True)
        
        # Bind events
        self.batch_tree.bind("<<TreeviewSelect>>", self._on_batch_selection_change)
        self.batch_tree.bind("<Double-1>", lambda e: self.view_transcript())
        self.batch_tree.bind("<Delete>", lambda e: self._delete_selected_transcripts())
        
        # Context menu
        self._build_batch_context_menu()
        
        # Initial load
        self.refresh_batch_list()
    
    def _build_batch_context_menu(self):
        """Right-click context menu for batch list."""
        self.batch_context_menu = tk.Menu(self.root, tearoff=0)
        self.batch_context_menu.add_command(label="👁 View Transcript", command=self.view_transcript)
        self.batch_context_menu.add_command(label="📤 Export Selected", command=self._export_batch_transcripts)
        self.batch_context_menu.add_separator()
        self.batch_context_menu.add_command(label="🗑 Delete Transcript", command=self._delete_selected_transcripts)
        self.batch_context_menu.add_command(label="🔄 Re-transcribe", command=self._retranscribe_selected)
        self.batch_context_menu.add_separator()
        self.batch_context_menu.add_command(label="📁 Open Audio Folder", command=self._open_audio_folder)
        self.batch_context_menu.add_command(label="📁 Open Transcripts Folder", command=self._open_transcripts_folder)
        
        def show_context_menu(event):
            item = self.batch_tree.identify_row(event.y)
            if item:
                self.batch_tree.selection_set(item)
                self.batch_context_menu.tk_popup(event.x_root, event.y_root)
        
        self.batch_tree.bind("<Button-3>", show_context_menu)
        self.batch_tree.bind("<Control-Button-1>", show_context_menu)
    
    def _build_speaker_tab(self, parent: ttk.Frame):
        """Speaker-labeled transcription tab."""
        # Source selection
        source_frame = ttk.LabelFrame(parent, text="Audio Source", padding=10)
        source_frame.pack(fill="x", pady=(0, 10))
        
        # Drag-drop area
        drop_frame = ttk.Frame(source_frame)
        drop_frame.pack(fill="x", pady=(0, 10))
        
        ttk.Label(drop_frame, text="📁 Drag audio file here or browse:").pack(side="left")
        
        self.speaker_source_var = tk.StringVar()
        self.speaker_source_combo = ttk.Combobox(
            drop_frame,
            textvariable=self.speaker_source_var,
            width=50,
        )
        self.speaker_source_combo.pack(side="left", padx=(10, 10), fill="x", expand=True)
        self.tooltip_manager.add(self.speaker_source_combo, "Select or drop an audio file")
        
        self.speaker_browse_button = ttk.Button(
            drop_frame, text="Browse…", command=self.browse_speaker_source
        )
        self.speaker_browse_button.pack(side="left")
        
        # Audio preview player
        self.speaker_preview_frame = ttk.Frame(source_frame)
        self.speaker_preview_frame.pack(fill="x", pady=(5, 0))
        self._build_audio_preview(self.speaker_preview_frame, self.speaker_source_var)
        
        # Action buttons
        action_frame = ttk.Frame(parent)
        action_frame.pack(fill="x", pady=(0, 10))
        
        self.transcribe_speaker_button = ttk.Button(
            action_frame, text="▶ Transcribe with Speaker Labels",
            command=self.start_speaker_transcription,
            bootstyle="success" if TTKBOOTSTRAP_AVAILABLE else None
        )
        self.transcribe_speaker_button.pack(side="left")
        self.tooltip_manager.add(self.transcribe_speaker_button, "Start speaker diarization (Ctrl+Shift+Enter)")
        
        self.cancel_speaker_button = ttk.Button(
            action_frame, text="⏹ Cancel", command=self._cancel_speaker_transcription,
            state="disabled", bootstyle="danger" if TTKBOOTSTRAP_AVAILABLE else None
        )
        self.cancel_speaker_button.pack(side="left", padx=(8, 0))
        
        self.save_speaker_button = ttk.Button(
            action_frame, text="💾 Save conversation.json", command=self.save_speaker_output,
            state="disabled"
        )
        self.save_speaker_button.pack(side="left", padx=(8, 0))
        self.tooltip_manager.add(self.save_speaker_button, "Save speaker-labeled transcript (Ctrl+S)")
        
        self.export_speaker_button = ttk.Button(
            action_frame, text="📤 Export", command=self._export_speaker_transcript,
            state="disabled"
        )
        self.export_speaker_button.pack(side="left", padx=(8, 0))
        self.tooltip_manager.add(self.export_speaker_button, "Export in different formats")
        
        self.speaker_status_label = ttk.Label(action_frame, text="Ready", foreground="gray")
        self.speaker_status_label.pack(side="left", padx=(12, 0))
        
        # Output area
        output_frame = ttk.LabelFrame(parent, text="Speaker-Labeled Transcript", padding=10)
        output_frame.pack(fill="both", expand=True)
        
        # Output toolbar
        out_toolbar = ttk.Frame(output_frame)
        out_toolbar.pack(fill="x", pady=(0, 8))
        
        ttk.Button(out_toolbar, text="📋 Copy All", command=self._copy_speaker_output,
                  bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="left", padx=2)
        ttk.Button(out_toolbar, text="🗑 Clear", command=self._clear_speaker_output,
                  bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="left", padx=2)
        
        self.speaker_output = scrolledtext.ScrolledText(
            output_frame, height=20, state="disabled", wrap="word",
            font=("Consolas", 10)
        )
        self.speaker_output.pack(fill="both", expand=True)
        
        # Configure tags for speaker coloring
        self.speaker_output.tag_configure("speaker_a", foreground="#4a90d9", font=("Consolas", 10, "bold"))
        self.speaker_output.tag_configure("speaker_b", foreground="#e67e22", font=("Consolas", 10, "bold"))
        self.speaker_output.tag_configure("speaker_c", foreground="#27ae60", font=("Consolas", 10, "bold"))
        self.speaker_output.tag_configure("speaker_d", foreground="#8e44ad", font=("Consolas", 10, "bold"))
        self.speaker_output.tag_configure("speaker_text", foreground="#ecf0f1")
        
        self._refresh_speaker_sources()
    
    def _build_audio_preview(self, parent: ttk.Frame, source_var: tk.StringVar):
        """Build audio preview player widget."""
        if not PYGAME_AVAILABLE:
            ttk.Label(parent, text="💡 Install pygame for audio preview: pip install pygame",
                     foreground="gray", font=("TkDefaultFont", 8)).pack(side="left")
            return
        
        self.preview_play_btn = ttk.Button(parent, text="▶ Play", width=8,
                                          command=lambda: self._play_preview(source_var.get()))
        self.preview_play_btn.pack(side="left", padx=(0, 5))
        self.tooltip_manager.add(self.preview_play_btn, "Play audio preview (Space)")
        
        self.preview_pause_btn = ttk.Button(parent, text="⏸ Pause", width=8,
                                           command=self.audio_player.pause, state="disabled")
        self.preview_pause_btn.pack(side="left", padx=5)
        
        self.preview_stop_btn = ttk.Button(parent, text="⏹ Stop", width=8,
                                          command=self.audio_player.stop, state="disabled")
        self.preview_stop_btn.pack(side="left", padx=5)
        
        # Volume slider
        ttk.Label(parent, text="🔊").pack(side="left", padx=(10, 2))
        self.preview_volume = tk.DoubleVar(value=0.7)
        volume_scale = ttk.Scale(parent, from_=0, to=1, variable=self.preview_volume,
                                orient="horizontal", length=100,
                                command=lambda v: self.audio_player.set_volume(float(v)))
        volume_scale.pack(side="left", padx=5)
        
        # Time label
        self.preview_time_label = ttk.Label(parent, text="0:00 / 0:00", foreground="gray")
        self.preview_time_label.pack(side="left", padx=10)
        
        # Update preview when source changes
        source_var.trace_add("write", lambda *_: self._on_preview_source_change(source_var.get()))
    
    def _build_history_tab(self, parent: ttk.Frame):
        """Transcription history tab."""
        # Toolbar
        hist_toolbar = ttk.Frame(parent)
        hist_toolbar.pack(fill="x", pady=(0, 8))
        
        ttk.Label(hist_toolbar, text="📜 Transcription History", font=("TkDefaultFont", 11, "bold")).pack(side="left")
        
        ttk.Button(hist_toolbar, text="🗑 Clear History", command=self._clear_history,
                  bootstyle="danger-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="right", padx=5)
        ttk.Button(hist_toolbar, text="📤 Export History", command=self._export_history,
                  bootstyle="secondary-outline" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="right", padx=5)
        
        # History tree frame
        hist_frame = ttk.Frame(parent)
        hist_frame.pack(fill="both", expand=True)

        # History tree
        hist_columns = ("time", "file", "type", "status", "duration", "language")
        self.history_tree = ttk.Treeview(hist_frame, columns=hist_columns, show="headings", height=20)
        
        self.history_tree.heading("time", text="🕐 Time", command=lambda: self._sort_history("time"))
        self.history_tree.heading("file", text="📄 File", command=lambda: self._sort_history("file"))
        self.history_tree.heading("type", text="Type", command=lambda: self._sort_history("type"))
        self.history_tree.heading("status", text="Status", command=lambda: self._sort_history("status"))
        self.history_tree.heading("duration", text="Duration", command=lambda: self._sort_history("duration"))
        self.history_tree.heading("language", text="Language", command=lambda: self._sort_history("language"))
        
        self.history_tree.column("time", width=150, anchor="center")
        self.history_tree.column("file", width=250, anchor="w", stretch=True)
        self.history_tree.column("type", width=100, anchor="center")
        self.history_tree.column("status", width=100, anchor="center")
        self.history_tree.column("duration", width=80, anchor="center")
        self.history_tree.column("language", width=120, anchor="center")
        
        hist_vsb = ttk.Scrollbar(hist_frame, orient="vertical", command=self.history_tree.yview)
        hist_hsb = ttk.Scrollbar(hist_frame, orient="horizontal", command=self.history_tree.xview)
        self.history_tree.configure(yscrollcommand=hist_vsb.set, xscrollcommand=hist_hsb.set)
        
        self.history_tree.grid(row=0, column=0, sticky="nsew")
        hist_vsb.grid(row=0, column=1, sticky="ns")
        hist_hsb.grid(row=1, column=0, sticky="ew")
        hist_frame.grid_rowconfigure(0, weight=1)
        hist_frame.grid_columnconfigure(0, weight=1)
        
        # Load existing history
        self._load_history()
        
        # Context menu
        self.history_context_menu = tk.Menu(self.root, tearoff=0)
        self.history_context_menu.add_command(label="👁 View Transcript", command=self._view_history_transcript)
        self.history_context_menu.add_command(label="📤 Export", command=self._export_history_item)
        self.history_context_menu.add_separator()
        self.history_context_menu.add_command(label="🗑 Remove from History", command=self._remove_history_item)
        
        def show_hist_menu(event):
            item = self.history_tree.identify_row(event.y)
            if item:
                self.history_tree.selection_set(item)
                self.history_context_menu.tk_popup(event.x_root, event.y_root)
        
        self.history_tree.bind("<Button-3>", show_hist_menu)
        self.history_tree.bind("<Double-1>", lambda e: self._view_history_transcript())
    
    def _build_status_bar(self):
        """Bottom status bar with info and shortcuts hint."""
        self.status_bar = ttk.Frame(self.main_container)
        self.status_bar.pack(fill="x", pady=(10, 0))
        
        # Left: API status
        self.api_status_label = ttk.Label(self.status_bar, text="🔑 API: Not configured", foreground="orange")
        self.api_status_label.pack(side="left")
        
        # Center: Shortcuts hint
        shortcuts_text = "Shortcuts: Ctrl+Enter=Transcribe | Ctrl+S=Save | Ctrl+E=Export | F5=Refresh | Esc=Cancel | Ctrl+=/=Zoom"
        self.shortcuts_label = ttk.Label(self.status_bar, text=shortcuts_text, foreground="gray", font=("TkDefaultFont", 8))
        self.shortcuts_label.pack(side="left", padx=20)
        
        # Right: Version/info
        self.version_label = ttk.Label(self.status_bar, text="v1.0 | AssemblyAI", foreground="gray")
        self.version_label.pack(side="right")
    
    # ─── Event Handlers & Callbacks ──────────────────────────────────────
    
    def _on_theme_change(self, event=None):
        """Handle theme change."""
        new_theme = self.theme_var.get()
        self.settings["theme"] = new_theme
        if TTKBOOTSTRAP_AVAILABLE:
            try:
                if new_theme in self.root.style.theme_names():
                    self.root.style.theme_use(new_theme)
            except Exception:
                pass
        messagebox.showinfo("Theme Changed", f"Theme set to '{new_theme}'.\nSome changes may require restart.", parent=self.root)
    
    def _on_language_change(self):
        """Handle language selection change."""
        self.settings["language"] = self.language_var.get()
    
    def _on_tab_change(self, event):
        """Handle tab change."""
        self._save_window_state()
        # Refresh speaker sources when switching to speaker tab
        if self.notebook.select() == str(self.speaker_frame):
            self._refresh_speaker_sources()
    
    def _bind_shortcuts(self):
        """Bind keyboard shortcuts."""
        shortcuts = {
            "<Control-Return>": lambda e: self.start_batch_transcription(),
            "<Control-Shift-Return>": lambda e: self.start_speaker_transcription(),
            "<Control-s>": lambda e: self.save_speaker_output(),
            "<Control-e>": lambda e: self._export_batch_transcripts(),
            "<Control-comma>": lambda e: self._open_settings(),
            "<F5>": lambda e: self.refresh_batch_list(),
            "<Escape>": lambda e: self._handle_escape(),
            "<Delete>": lambda e: self._delete_selected_transcripts(),
            "<Control-a>": lambda e: self._select_all_batch(),
            "<Control-f>": lambda e: self._focus_search(),
            "<Control-plus>": lambda e: self._adjust_font_size(1),
            "<Control-minus>": lambda e: self._adjust_font_size(-1),
            "<Control-0>": lambda e: self._reset_font_size(),
            "<space>": lambda e: self._handle_space_preview(),
        }
        
        for key, cmd in shortcuts.items():
            self.root.bind(key, cmd)
    
    def _handle_escape(self):
        """Handle Escape key - cancel current operation."""
        if self.busy:
            if self.batch_worker_thread and self.batch_worker_thread.is_alive():
                self._cancel_batch_transcription()
            elif self.speaker_worker_thread and self.speaker_worker_thread.is_alive():
                self._cancel_speaker_transcription()
    
    def _handle_space_preview(self):
        """Space to play/pause preview if focused on speaker tab."""
        if self.notebook.select() == str(self.speaker_frame):
            if self.audio_player.is_playing():
                self.audio_player.pause()
            else:
                self._play_preview(self.speaker_source_var.get())
    
    def _focus_search(self):
        """Focus the batch search entry."""
        for child in self.batch_frame.winfo_children():
            if isinstance(child, ttk.Frame):
                for w in child.winfo_children():
                    if isinstance(w, ttk.Entry):
                        w.focus_set()
                        return
    
    def _adjust_font_size(self, delta: int):
        """Adjust font size."""
        current = self.settings.get("font_size", 10)
        new_size = max(8, min(18, current + delta))
        self.settings["font_size"] = new_size
        self.root.option_add("*Font", f"TkDefaultFont {new_size}")
        messagebox.showinfo("Font Size", f"Font size set to {new_size}. Restart for full effect.", parent=self.root)
    
    def _reset_font_size(self):
        """Reset font size to default."""
        self.settings["font_size"] = 10
        self.root.option_add("*Font", "TkDefaultFont 10")
        messagebox.showinfo("Font Size", "Font size reset to default. Restart for full effect.", parent=self.root)
    
    def _setup_drag_drop(self):
        """Setup drag and drop for audio files."""
        if not DND_AVAILABLE:
            return
        
        try:
            # Make speaker combo a drop target
            self.speaker_source_combo.drop_target_register(DND_FILES)
            self.speaker_source_combo.dnd_bind('<<Drop>>', self._on_drop_audio)
            
            # Also make batch tree a drop target
            self.batch_tree.drop_target_register(DND_FILES)
            self.batch_tree.dnd_bind('<<Drop>>', self._on_drop_audio_batch)
        except Exception:
            pass
    
    def _on_drop_audio(self, event):
        """Handle dropped audio file on speaker tab."""
        files = self._parse_drop_files(event.data)
        audio_files = [f for f in files if self._is_audio_file(f)]
        if audio_files:
            self.speaker_source_var.set(audio_files[0])
    
    def _on_drop_audio_batch(self, event):
        """Handle dropped audio files on batch tab - copy to audio folder."""
        files = self._parse_drop_files(event.data)
        audio_files = [f for f in files if self._is_audio_file(f)]
        if audio_files:
            self._copy_files_to_audio_dir(audio_files)
            self.refresh_batch_list()
    
    def _parse_drop_files(self, data: str) -> list[str]:
        """Parse dropped file paths from tkinterdnd2 event data."""
        # Data format varies by platform
        files = []
        for part in data.split():
            # Remove braces on Windows
            part = part.strip('{}')
            if os.path.exists(part):
                files.append(part)
        return files
    
    def _is_audio_file(self, path: str) -> bool:
        """Check if file has supported audio extension."""
        ext = Path(path).suffix.lower()
        return ext in t.SUPPORTED_EXTENSIONS
    
    def _copy_files_to_audio_dir(self, files: list[str]):
        """Copy dropped files to audio directory."""
        import shutil
        t.ensure_directories()
        for f in files:
            try:
                dest = Path(t.AUDIO_DIR) / Path(f).name
                shutil.copy2(f, dest)
            except Exception as e:
                self._append_batch_log(f"Failed to copy {f}: {e}")
    
    # ─── API Key Handling ────────────────────────────────────────────────
    
    def _try_load_api_key_from_env(self):
        if t.get_api_key():
            try:
                t.init_assemblyai()
                self.api_key_configured = True
            except ValueError:
                self.api_key_configured = False
    
    def _submit_api_key(self):
        api_key = self.api_key_text.get("1.0", "end").strip()
        if not api_key:
            messagebox.showerror("API Key Missing", "Please enter your AssemblyAI API key.", parent=self.root)
            self.api_key_text.focus_set()
            return
        
        try:
            t.configure_assemblyai(api_key)
            if self.save_env_var.get():
                t.save_api_key_to_env(api_key)
        except (ValueError, OSError) as exc:
            messagebox.showerror("Invalid API Key", str(exc), parent=self.root)
            return
        
        self.api_key_configured = True
        self.api_key_text.delete("1.0", "end")
        self.api_key_frame.pack_forget()
        self._update_app_state()
        self._update_api_status()
        messagebox.showinfo("API Key Saved", "Your API key is set. You can start transcribing.", parent=self.root)
    
    def _update_api_status(self):
        """Update API status label in status bar."""
        if self.api_key_configured:
            self.api_status_label.configure(text="🔑 API: Connected", foreground="green")
        else:
            self.api_status_label.configure(text="🔑 API: Not configured", foreground="orange")
    
    def _require_api_key(self) -> bool:
        if self.api_key_configured:
            return True
        messagebox.showwarning("API Key Required", 
            "Enter your AssemblyAI API key in the box at the top of the window first.",
            parent=self.root)
        if hasattr(self, 'api_key_text'):
            self.api_key_text.focus_set()
        return False
    
    # ─── App State Management ────────────────────────────────────────────
    
    def _update_app_state(self):
        """Enable/disable UI elements based on state."""
        if self.api_key_configured:
            self.api_key_frame.pack_forget()
        elif not self.api_key_frame.winfo_ismapped():
            # Insert before toolbar
            children = list(self.main_container.winfo_children())
            toolbar_idx = children.index(self.toolbar) if self.toolbar in children else 1
            self.api_key_frame.pack(fill="x", pady=(0, 10), before=self.toolbar)
            self.api_key_text.focus_set()
        
        enabled = self.api_key_configured and not self.busy
        readonly_state = "readonly" if enabled else "disabled"
        state = "normal" if enabled else "disabled"
        
        # Language and batch controls
        self.language_combo.configure(state=readonly_state)
        self.refresh_button.configure(state=state)
        self.transcribe_batch_button.configure(state=state)
        self.cancel_batch_button.configure(state="normal" if (self.busy and self.batch_worker_thread) else "disabled")
        self.view_transcript_button.configure(state="disabled")  # Updated on selection
        self.export_batch_button.configure(state="disabled")
        self.speaker_browse_button.configure(state=state)
        self.speaker_source_combo.configure(state=state)
        self.transcribe_speaker_button.configure(state=state)
        self.cancel_speaker_button.configure(state="normal" if (self.busy and self.speaker_worker_thread) else "disabled")
        
        # Speaker output controls
        if not self.api_key_configured:
            self.save_speaker_button.configure(state="disabled")
            self.export_speaker_button.configure(state="disabled")
        elif self.busy:
            self.save_speaker_button.configure(state="disabled")
            self.export_speaker_button.configure(state="disabled")
        elif self.speaker_conversation:
            self.save_speaker_button.configure(state="normal")
            self.export_speaker_button.configure(state="normal")
        else:
            self.save_speaker_button.configure(state="disabled")
            self.export_speaker_button.configure(state="disabled")
    
    def _set_busy(self, busy: bool):
        self.busy = busy
        self.cancellation_token.reset()
        self._update_app_state()
    
    # ─── Batch Transcription ─────────────────────────────────────────────
    
    def refresh_batch_list(self):
        """Refresh the batch file list."""
        self.batch_tree.delete(*self.batch_tree.get_children())
        self.batch_tree_items.clear()
        
        for audio_name in t.list_audio_files():
            status = t.get_file_status(audio_name)
            display_status = "✅ Done" if status == "done" else "⏳ Pending"
            
            # Get file info
            audio_path = Path(t.AUDIO_DIR) / audio_name
            try:
                stat = audio_path.stat()
                size_str = self._format_size(stat.st_size)
                mod_str = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M")
            except Exception:
                size_str = "—"
                mod_str = "—"
            
            # Try to get duration from transcript if exists
            duration_str = "—"
            transcript_path = Path(t.transcript_path_for(audio_name))
            if transcript_path.exists():
                try:
                    with open(transcript_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        # Duration might be in metadata
                        if "audio_duration" in data:
                            duration_str = self._format_duration(data["audio_duration"])
                except Exception:
                    pass
            
            item_id = self.batch_tree.insert("", "end", values=(
                audio_name, display_status, size_str, mod_str, duration_str
            ))
            self.batch_tree_items[item_id] = audio_name
        
        # Update status label
        pending = len(t.get_pending_files())
        total = len(t.list_audio_files())
        if pending:
            self.batch_status_label.configure(text=f"{pending} of {total} files pending")
        else:
            self.batch_status_label.configure(text=f"All {total} files transcribed")
        
        self._filter_batch_list()
    
    def _filter_batch_list(self):
        """Filter batch list based on search text."""
        search = self.batch_search_var.get().lower()
        # Use get_children("") to get ALL items including detached ones
        for item in self.batch_tree.get_children(""):
            values = self.batch_tree.item(item, "values")
            if not search or any(search in str(v).lower() for v in values):
                self.batch_tree.reattach(item, "", "end")
            else:
                self.batch_tree.detach(item)
    
    def _sort_tree(self, col: str):
        """Sort tree by column."""
        items = [(self.batch_tree.set(item, col), item) for item in self.batch_tree.get_children("")]
        # Try numeric sort for size/duration
        try:
            items.sort(key=lambda x: float(x[0].replace(",", "").replace("MB", "").replace("KB", "").strip()))
        except Exception:
            items.sort(key=lambda x: x[0].lower())
        
        for index, (_, item) in enumerate(items):
            self.batch_tree.move(item, "", index)
    
    def _on_batch_selection_change(self, event=None):
        """Handle batch tree selection change."""
        selection = self.batch_tree.selection()
        if not selection:
            self.view_transcript_button.configure(state="disabled")
            self.export_batch_button.configure(state="disabled")
            return
        
        # Check if any selected has transcript
        has_transcript = False
        for item in selection:
            audio_name = self.batch_tree_items.get(item)
            if audio_name and os.path.exists(t.transcript_path_for(audio_name)):
                has_transcript = True
                break
        
        self.view_transcript_button.configure(state="normal" if has_transcript else "disabled")
        self.export_batch_button.configure(state="normal" if has_transcript else "disabled")
    
    def _select_all_batch(self):
        """Select all visible items in batch tree."""
        for item in self.batch_tree.get_children(""):
            self.batch_tree.selection_add(item)
    
    def start_batch_transcription(self):
        """Start batch transcription."""
        if self.busy or not self._require_api_key():
            return
        
        try:
            language_code = self.get_selected_language_code()
        except ValueError as exc:
            messagebox.showerror("Language Error", str(exc), parent=self.root)
            return
        
        pending_files = t.get_pending_files()
        if not pending_files:
            messagebox.showinfo("Batch Transcription", "All audio files already have transcripts.", parent=self.root)
            return
        
        # Confirm if many files
        if len(pending_files) > 10:
            if not messagebox.askyesno("Confirm", f"Transcribe {len(pending_files)} files? This may take a while.", parent=self.root):
                return
        
        self._set_busy(True)
        self.batch_progress.configure(maximum=len(pending_files), value=0)
        self._append_batch_log(f"Starting batch transcription ({len(pending_files)} file(s))...")
        self.transcription_start_time = time.time()
        
        def worker():
            try:
                def on_progress(current: int, total: int, name: str, status: str):
                    if self.cancellation_token.is_cancelled:
                        raise InterruptedError("Cancelled by user")
                    self.event_queue.put(("batch_progress", current, total, name, status))
                
                results = t.transcribe_batch(
                    language_code=language_code,
                    on_progress=on_progress,
                )
                if not self.cancellation_token.is_cancelled:
                    self.event_queue.put(("batch_done", results))
            except InterruptedError:
                self.event_queue.put(("batch_cancelled",))
            except Exception as exc:
                self.event_queue.put(("error", "Batch Transcription", str(exc)))
        
        self.batch_worker_thread = threading.Thread(target=worker, daemon=True)
        self.batch_worker_thread.start()
    
    def _cancel_batch_transcription(self):
        """Cancel batch transcription."""
        if self.busy and self.batch_worker_thread and self.batch_worker_thread.is_alive():
            self.cancellation_token.cancel()
            self._append_batch_log("⏹ Cancellation requested...")
            self.cancel_batch_button.configure(state="disabled")
    
    def _handle_event(self, event: tuple):
        kind = event[0]
        
        if kind == "batch_progress":
            _, current, total, name, status = event
            if status == "transcribing" and self.transcription_start_time is None:
                self.transcription_start_time = time.time()
                self.current_transcription_file = name
            
            self.batch_progress.configure(maximum=total, value=current - 1 if status == "transcribing" else current)
            
            if status == "transcribing":
                self.batch_current_label.configure(text=f"Transcribing {current}/{total}: {name}")
                self._append_batch_log(f"🎙 Transcribing ({current}/{total}): {t.AUDIO_DIR}/{name}")
            elif status == "done":
                self.batch_current_label.configure(text=f"✅ Completed {current}/{total}")
            elif status == "failed":
                self.batch_current_label.configure(text=f"❌ Failed {current}/{total}")
        
        elif kind == "batch_done":
            _, results = event
            success_count = sum(1 for _, success, _ in results if success)
            fail_count = len(results) - success_count
            
            for audio_name, success, message in results:
                if success:
                    self._append_batch_log(f"✅ {message}")
                else:
                    self._append_batch_log(f"❌ Failed: {audio_name} → {message}")
            
            self.batch_progress.configure(value=self.batch_progress["maximum"])
            self.batch_current_label.configure(text=f"Complete: {success_count} succeeded, {fail_count} failed")
            self.transcription_start_time = None
            self.current_transcription_file = None
            self.refresh_batch_list()
            self._add_to_history("batch", results)
            self._set_busy(False)
        
        elif kind == "batch_cancelled":
            self._append_batch_log("⏹ Transcription cancelled by user")
            self.batch_current_label.configure(text="Cancelled")
            self.transcription_start_time = None
            self._set_busy(False)
        
        elif kind == "speaker_done":
            _, result = event
            if not isinstance(result, list):
                messagebox.showerror("Speaker Transcription", "Unexpected result format.", parent=self.root)
                self._set_busy(False)
                return
            
            self.speaker_conversation = result
            self._display_speaker_conversation(result)
            self.speaker_status_label.configure(text=f"✅ Done ({len(result)} utterances)")
            
            # Auto-save if enabled
            if self.settings.get("auto_save_speaker", True):
                path = t.save_speaker_conversation(result)
                self.speaker_status_label.configure(text=f"✅ Done — auto-saved to {path}")
            
            self.save_speaker_button.configure(state="normal")
            self.export_speaker_button.configure(state="normal")
            self._add_to_history("speaker", result)
            self._set_busy(False)
        
        elif kind == "error":
            _, title, message = event
            messagebox.showerror(title, message, parent=self.root)
            self.batch_current_label.configure(text="Error")
            self.speaker_status_label.configure(text="Error")
            self._set_busy(False)
    
    def _append_batch_log(self, message: str):
        self.batch_log.configure(state="normal")
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.batch_log.insert("end", f"[{timestamp}] {message}\n")
        self.batch_log.see("end")
        self.batch_log.configure(state="disabled")
    
    def _clear_batch_log(self):
        self.batch_log.configure(state="normal")
        self.batch_log.delete("1.0", "end")
        self.batch_log.configure(state="disabled")
    
    def _copy_batch_log(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.batch_log.get("1.0", "end-1c"))
        self._append_batch_log("📋 Log copied to clipboard")
    
    def _save_batch_log(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="Save Transcription Log"
        )
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.batch_log.get("1.0", "end-1c"))
                self._append_batch_log(f"💾 Log saved to {path}")
            except Exception as e:
                messagebox.showerror("Save Failed", str(e), parent=self.root)
    
    def view_transcript(self):
        """View transcript of selected file."""
        selection = self.batch_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        audio_name = self.batch_tree_items.get(item)
        if not audio_name:
            return
        
        transcript_path = Path(t.transcript_path_for(audio_name))
        if not transcript_path.exists():
            messagebox.showinfo("View Transcript", f"No transcript found for {audio_name}", parent=self.root)
            return
        
        try:
            with open(transcript_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            transcript_text = data.get("transcript", "")
        except Exception as exc:
            messagebox.showerror("View Transcript", f"Failed to read transcript: {exc}", parent=self.root)
            return
        
        self._show_transcript_dialog(audio_name, transcript_text, data)
    
    def _show_transcript_dialog(self, title: str, text: str, metadata: dict = None):
        """Show transcript in a dialog with copy/export options."""
        dialog = tk.Toplevel(self.root)
        dialog.title(f"Transcript: {title}")
        dialog.geometry("800x600")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Toolbar
        toolbar = ttk.Frame(dialog, padding=10)
        toolbar.pack(fill="x")
        
        ttk.Button(toolbar, text="📋 Copy Text", command=lambda: self._copy_to_clipboard(text, dialog)).pack(side="left", padx=2)
        ttk.Button(toolbar, text="📤 Export", command=lambda: self._export_transcript_dialog(text, metadata, dialog)).pack(side="left", padx=2)
        ttk.Button(toolbar, text="🔍 Search", command=lambda: self._show_search_dialog(dialog, text_widget)).pack(side="left", padx=2)
        ttk.Button(toolbar, text="Close", command=dialog.destroy).pack(side="right", padx=2)
        
        # Info bar
        if metadata:
            info_frame = ttk.Frame(dialog, padding=(10, 0))
            info_frame.pack(fill="x")
            info_text = []
            if "language_code" in metadata:
                info_text.append(f"Language: {metadata['language_code']}")
            if "confidence" in metadata:
                info_text.append(f"Confidence: {metadata['confidence']:.1%}")
            if "audio_duration" in metadata:
                info_text.append(f"Duration: {self._format_duration(metadata['audio_duration'])}")
            ttk.Label(info_frame, text=" | ".join(info_text), foreground="gray").pack(side="left")
        
        # Text area
        text_frame = ttk.Frame(dialog)
        text_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        text_widget = scrolledtext.ScrolledText(text_frame, wrap="word", state="normal", font=("Consolas", 10))
        text_widget.pack(fill="both", expand=True)
        text_widget.insert("1.0", text)
        text_widget.configure(state="disabled")
        
        dialog.focus_set()
    
    def _export_batch_transcripts(self):
        """Export selected transcripts."""
        selection = self.batch_tree.selection()
        if not selection:
            return
        
        transcripts = []
        for item in selection:
            audio_name = self.batch_tree_items.get(item)
            if audio_name:
                path = Path(t.transcript_path_for(audio_name))
                if path.exists():
                    try:
                        with open(path, "r", encoding="utf-8") as f:
                            data = json.load(f)
                        transcripts.append((audio_name, data.get("transcript", ""), data))
                    except Exception:
                        pass
        
        if not transcripts:
            return
        
        self._export_multiple_transcripts(transcripts)
    
    def _export_multiple_transcripts(self, transcripts: list[tuple[str, str, dict]]):
        """Export multiple transcripts with format selection."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Export Transcripts")
        dialog.geometry("400x300")
        dialog.transient(self.root)
        dialog.grab_set()
        
        ttk.Label(dialog, text=f"Export {len(transcripts)} transcript(s)", font=("TkDefaultFont", 10, "bold")).pack(pady=10)
        
        format_var = tk.StringVar(value=self.settings.get("default_export_format", "txt"))
        ttk.Label(dialog, text="Format:").pack(anchor="w", padx=20)
        format_combo = ttk.Combobox(dialog, textvariable=format_var, 
                                   values=[f[1] for f in EXPORT_FORMATS], state="readonly")
        format_combo.pack(fill="x", padx=20, pady=5)
        
        # Options
        opt_frame = ttk.LabelFrame(dialog, text="Options", padding=10)
        opt_frame.pack(fill="x", padx=20, pady=10)
        
        separate_files = tk.BooleanVar(value=True)
        ttk.Checkbutton(opt_frame, text="Separate file per transcript", variable=separate_files).pack(anchor="w")
        
        include_metadata = tk.BooleanVar(value=True)
        ttk.Checkbutton(opt_frame, text="Include metadata (JSON only)", variable=include_metadata).pack(anchor="w")
        
        def do_export():
            fmt = format_var.get()
            self.settings["default_export_format"] = fmt
            
            if separate_files.get():
                dir_path = filedialog.askdirectory(title="Select Export Directory")
                if not dir_path:
                    return
                for audio_name, text, meta in transcripts:
                    base = Path(audio_name).stem
                    out_path = Path(dir_path) / f"{base}.{fmt}"
                    self._write_export(out_path, text, meta, fmt, include_metadata.get())
                messagebox.showinfo("Export Complete", f"Exported {len(transcripts)} files to {dir_path}", parent=dialog)
            else:
                path = filedialog.asksaveasfilename(
                    defaultextension=f".{fmt}",
                    filetypes=[(f"{f[0]}", f"*.{f[1]}") for f in EXPORT_FORMATS],
                    title="Save Combined Transcript"
                )
                if not path:
                    return
                combined = "\n\n---\n\n".join(f"[{name}]\n{text}" for name, text, _ in transcripts)
                self._write_export(Path(path), combined, {}, fmt, include_metadata.get())
                messagebox.showinfo("Export Complete", f"Saved to {path}", parent=dialog)
            
            dialog.destroy()
        
        ttk.Button(dialog, text="Export", command=do_export, bootstyle="success" if TTKBOOTSTRAP_AVAILABLE else None).pack(pady=10)
    
    def _write_export(self, path: Path, text: str, metadata: dict, fmt: str, include_meta: bool):
        """Write transcript in specified format."""
        if fmt == "txt":
            with open(path, "w", encoding="utf-8") as f:
                if include_meta and metadata:
                    f.write(f"# {metadata.get('file_name', path.stem)}\n")
                    f.write(f"# Language: {metadata.get('language_code', 'unknown')}\n")
                    f.write(f"# Generated: {datetime.now().isoformat()}\n\n")
                f.write(text)
        
        elif fmt == "srt":
            # Simple SRT - split by sentences (approximate)
            with open(path, "w", encoding="utf-8") as f:
                sentences = text.replace("\n", " ").split(". ")
                for i, sent in enumerate(sentences, 1):
                    if sent.strip():
                        f.write(f"{i}\n")
                        f.write(f"00:00:00,000 --> 00:00:05,000\n")  # Placeholder timing
                        f.write(f"{sent.strip()}\n\n")
        
        elif fmt == "vtt":
            with open(path, "w", encoding="utf-8") as f:
                f.write("WEBVTT\n\n")
                sentences = text.replace("\n", " ").split(". ")
                for i, sent in enumerate(sentences, 1):
                    if sent.strip():
                        f.write(f"{i}\n")
                        f.write(f"00:00:00.000 --> 00:00:05.000\n")
                        f.write(f"{sent.strip()}\n\n")
        
        elif fmt == "json":
            with open(path, "w", encoding="utf-8") as f:
                data = {"transcript": text}
                if include_meta:
                    data.update(metadata)
                json.dump(data, f, ensure_ascii=False, indent=2)
        
        elif fmt == "docx" and DOCX_AVAILABLE:
            doc = docx.Document()
            if include_meta and metadata:
                doc.add_heading(metadata.get('file_name', path.stem), 0)
                for k, v in metadata.items():
                    if k != 'transcript':
                        doc.add_paragraph(f"{k}: {v}")
                doc.add_paragraph("")
            doc.add_paragraph(text)
            doc.save(path)
    
    def _export_transcript_dialog(self, text: str, metadata: dict, parent_dialog):
        """Export single transcript from view dialog."""
        parent_dialog.destroy()
        self._export_multiple_transcripts([("transcript", text, metadata or {})])
    
    def _delete_selected_transcripts(self):
        """Delete transcript files for selected items."""
        selection = self.batch_tree.selection()
        if not selection:
            return
        
        files_to_delete = []
        for item in selection:
            audio_name = self.batch_tree_items.get(item)
            if audio_name:
                path = Path(t.transcript_path_for(audio_name))
                if path.exists():
                    files_to_delete.append((audio_name, path))
        
        if not files_to_delete:
            return
        
        if messagebox.askyesno("Delete Transcripts", 
                              f"Delete {len(files_to_delete)} transcript file(s)?\nAudio files will NOT be deleted.",
                              parent=self.root):
            for audio_name, path in files_to_delete:
                try:
                    path.unlink()
                    self._append_batch_log(f"🗑 Deleted transcript for {audio_name}")
                except Exception as e:
                    self._append_batch_log(f"❌ Failed to delete {audio_name}: {e}")
            self.refresh_batch_list()
    
    def _retranscribe_selected(self):
        """Re-transcribe selected files (delete existing transcripts first)."""
        selection = self.batch_tree.selection()
        if not selection:
            return
        
        audio_names = [self.batch_tree_items.get(item) for item in selection if self.batch_tree_items.get(item)]
        if not audio_names:
            return
        
        if messagebox.askyesno("Re-transcribe", 
                              f"Delete existing transcripts and re-transcribe {len(audio_names)} file(s)?",
                              parent=self.root):
            for name in audio_names:
                path = Path(t.transcript_path_for(name))
                if path.exists():
                    try:
                        path.unlink()
                    except Exception:
                        pass
            self.refresh_batch_list()
            self.start_batch_transcription()
    
    def _open_audio_folder(self):
        """Open audio folder in file manager."""
        t.ensure_directories()
        self._open_folder(t.AUDIO_DIR)
    
    def _open_transcripts_folder(self):
        """Open transcripts folder in file manager."""
        t.ensure_directories()
        self._open_folder(t.TRANSCRIPTS_DIR)
    
    def _open_folder(self, path: str):
        """Open folder in system file manager."""
        import subprocess
        import sys
        try:
            if sys.platform == "darwin":
                subprocess.run(["open", path])
            elif sys.platform == "win32":
                subprocess.run(["explorer", path])
            else:
                subprocess.run(["xdg-open", path])
        except Exception:
            pass
    
    # ─── Speaker Transcription ──────────────────────────────────────────
    
    def _refresh_speaker_sources(self):
        """Refresh speaker source dropdown."""
        values = [f"{t.AUDIO_DIR}/{name}" for name in t.list_audio_files()]
        self.speaker_source_combo.configure(values=values)
        if values and not self.speaker_source_var.get():
            self.speaker_source_var.set(values[0])
    
    def browse_speaker_source(self):
        """Browse for audio file."""
        path = filedialog.askopenfilename(
            title="Select Audio File",
            initialdir=os.path.abspath(t.AUDIO_DIR),
            filetypes=[
                ("Audio Files", "*.mp3 *.m4a *.wav *.mp4 *.aac *.flac *.ogg *.webm"),
                ("All Files", "*.*"),
            ],
        )
        if path:
            self.speaker_source_var.set(path)
    
    def _on_preview_source_change(self, source: str):
        """Handle preview source change."""
        # Reset player buttons
        self.preview_play_btn.configure(state="normal" if source else "disabled")
        self.preview_pause_btn.configure(state="disabled")
        self.preview_stop_btn.configure(state="disabled")
        self.preview_time_label.configure(text="0:00 / 0:00")
    
    def _play_preview(self, source: str):
        """Play audio preview."""
        if not source:
            return
        
        resolved = t.resolve_audio_path(source)
        if resolved.startswith(("http://", "https://")):
            messagebox.showinfo("Preview", "Cannot preview remote URLs. Download first.", parent=self.root)
            return
        
        if not os.path.exists(resolved):
            messagebox.showerror("Preview", f"File not found: {resolved}", parent=self.root)
            return
        
        if self.audio_player.play(resolved):
            self.preview_play_btn.configure(state="disabled")
            self.preview_pause_btn.configure(state="normal")
            self.preview_stop_btn.configure(state="normal")
            self._update_preview_time()
        else:
            messagebox.showerror("Preview", "Failed to play audio. Install pygame: pip install pygame", parent=self.root)
    
    def _update_preview_time(self):
        """Update preview time label."""
        if self.audio_player.is_playing():
            # Note: pygame doesn't easily give current position without more work
            self.preview_time_label.configure(text="▶ Playing...")
            self.root.after(1000, self._update_preview_time)
        else:
            self.preview_play_btn.configure(state="normal")
            self.preview_pause_btn.configure(state="disabled")
            self.preview_stop_btn.configure(state="disabled")
            self.preview_time_label.configure(text="0:00 / 0:00")
    
    def start_speaker_transcription(self):
        """Start speaker-labeled transcription."""
        if self.busy or not self._require_api_key():
            return
        
        source = self.speaker_source_var.get().strip()
        if not source:
            messagebox.showerror("Speaker Transcription", "Select or enter an audio source.", parent=self.root)
            return
        
        source = t.resolve_audio_path(source)
        if not source.startswith(("http://", "https://")) and not os.path.exists(source):
            messagebox.showerror("Speaker Transcription", 
                f"Audio file not found.\nPut files in {t.AUDIO_DIR}/ or choose a valid path/URL.",
                parent=self.root)
            return
        
        try:
            language_code = self.get_selected_language_code()
        except ValueError as exc:
            messagebox.showerror("Language Error", str(exc), parent=self.root)
            return
        
        self._set_busy(True)
        self.speaker_status_label.configure(text="🎙 Transcribing with speaker labels…")
        self._set_speaker_output("")
        self.transcription_start_time = time.time()
        
        def worker():
            try:
                result = t.transcribe_file(
                    source,
                    speaker_labels=True,
                    language_code=language_code,
                )
                if not self.cancellation_token.is_cancelled:
                    self.event_queue.put(("speaker_done", result))
            except InterruptedError:
                self.event_queue.put(("speaker_cancelled",))
            except Exception as exc:
                self.event_queue.put(("error", "Speaker Transcription", str(exc)))
        
        self.speaker_worker_thread = threading.Thread(target=worker, daemon=True)
        self.speaker_worker_thread.start()
    
    def _cancel_speaker_transcription(self):
        """Cancel speaker transcription."""
        if self.busy and self.speaker_worker_thread and self.speaker_worker_thread.is_alive():
            self.cancellation_token.cancel()
            self.speaker_status_label.configure(text="⏹ Cancelling...")
            self.cancel_speaker_button.configure(state="disabled")
    
    def _display_speaker_conversation(self, conversation: list[dict]):
        """Display speaker conversation with color coding."""
        self.speaker_output.configure(state="normal")
        self.speaker_output.delete("1.0", "end")
        
        speaker_colors = ["speaker_a", "speaker_b", "speaker_c", "speaker_d"]
        speaker_map = {}  # Map speaker ID to color index
        
        for entry in conversation:
            speaker = entry.get("speaker", "A")
            text = entry.get("text", "")
            
            if speaker not in speaker_map:
                speaker_map[speaker] = len(speaker_map) % len(speaker_colors)
            
            color_tag = speaker_colors[speaker_map[speaker]]
            
            self.speaker_output.insert("end", f"Speaker {speaker}: ", color_tag)
            self.speaker_output.insert("end", f"{text}\n\n", "speaker_text")
        
        self.speaker_output.configure(state="disabled")
        self.speaker_output.see("1.0")
    
    def _set_speaker_output(self, text: str):
        self.speaker_output.configure(state="normal")
        self.speaker_output.delete("1.0", "end")
        if text:
            self.speaker_output.insert("1.0", text)
        self.speaker_output.configure(state="disabled")
    
    def _copy_speaker_output(self):
        text = self.speaker_output.get("1.0", "end-1c")
        if text.strip():
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
            self.speaker_status_label.configure(text="📋 Copied to clipboard")
        else:
            self.speaker_status_label.configure(text="Nothing to copy")
    
    def _clear_speaker_output(self):
        self._set_speaker_output("")
        self.speaker_conversation = None
        self.save_speaker_button.configure(state="disabled")
        self.export_speaker_button.configure(state="disabled")
        self.speaker_status_label.configure(text="Ready")
    
    def save_speaker_output(self):
        if not self._require_api_key():
            return
        
        if not self.speaker_conversation:
            messagebox.showinfo("Speaker Transcription", "Nothing to save. Transcribe a file first.", parent=self.root)
            return
        
        path = t.save_speaker_conversation(self.speaker_conversation)
        self.speaker_status_label.configure(text=f"💾 Saved to {path}")
        messagebox.showinfo("Saved", f"Conversation saved to {path}", parent=self.root)
    
    def _export_speaker_transcript(self):
        if not self.speaker_conversation:
            return
        
        # Format as readable text
        lines = [f"Speaker {e['speaker']}: {e['text']}" for e in self.speaker_conversation]
        text = "\n".join(lines)
        
        self._export_multiple_transcripts([("conversation", text, {"type": "speaker_labeled"})])
    
    # ─── History Tab ─────────────────────────────────────────────────────
    
    def _load_history(self):
        """Load transcription history from file."""
        history_file = Path(t.TRANSCRIPTS_DIR) / ".history.json"
        if history_file.exists():
            try:
                with open(history_file, "r", encoding="utf-8") as f:
                    self.history_items = json.load(f)
            except Exception:
                self.history_items = []
        else:
            self.history_items = []
        
        self._populate_history_tree()
    
    def _save_history(self):
        """Save history to file."""
        history_file = Path(t.TRANSCRIPTS_DIR) / ".history.json"
        try:
            with open(history_file, "w", encoding="utf-8") as f:
                json.dump(self.history_items[-100:], f, ensure_ascii=False, indent=2)  # Keep last 100
        except Exception:
            pass
    
    def _add_to_history(self, trans_type: str, result):
        """Add transcription to history."""
        timestamp = datetime.now().isoformat()
        
        if trans_type == "batch":
            for audio_name, success, message in result:
                self.history_items.append({
                    "time": timestamp,
                    "file": audio_name,
                    "type": "Batch",
                    "status": "Success" if success else "Failed",
                    "duration": 0,  # Could track actual duration
                    "language": self.language_var.get(),
                })
        else:  # speaker
            self.history_items.append({
                "time": timestamp,
                "file": self.speaker_source_var.get(),
                "type": "Speaker",
                "status": "Success",
                "duration": 0,
                "language": self.language_var.get(),
            })
        
        # Keep only last 100
        self.history_items = self.history_items[-100:]
        self._save_history()
        self._populate_history_tree()
    
    def _populate_history_tree(self):
        """Populate history tree with items."""
        self.history_tree.delete(*self.history_tree.get_children())
        
        for item in reversed(self.history_items):  # Most recent first
            time_str = datetime.fromisoformat(item["time"]).strftime("%Y-%m-%d %H:%M:%S")
            status_icon = "✅" if item["status"] == "Success" else "❌"
            
            self.history_tree.insert("", "end", values=(
                time_str,
                item["file"],
                item["type"],
                f"{status_icon} {item['status']}",
                self._format_duration(item.get("duration", 0)),
                item.get("language", "Auto"),
            ))
    
    def _sort_history(self, col: str):
        """Sort history tree."""
        items = [(self.history_tree.set(item, col), item) for item in self.history_tree.get_children("")]
        try:
            items.sort(key=lambda x: float(x[0].replace(":", "").replace(".", "")))
        except Exception:
            items.sort(key=lambda x: x[0].lower())
        
        for index, (_, item) in enumerate(items):
            self.history_tree.move(item, "", index)
    
    def _view_history_transcript(self):
        """View transcript from history."""
        selection = self.history_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        values = self.history_tree.item(item, "values")
        filename = values[1]
        
        # Find transcript
        if values[2] == "Batch":
            transcript_path = Path(t.transcript_path_for(filename))
        else:
            transcript_path = Path("conversation.json")
        
        if transcript_path.exists():
            try:
                with open(transcript_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, list):  # Speaker format
                    text = "\n".join(f"Speaker {e['speaker']}: {e['text']}" for e in data)
                else:
                    text = data.get("transcript", "")
                self._show_transcript_dialog(filename, text, data)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load transcript: {e}", parent=self.root)
        else:
            messagebox.showinfo("Not Found", "Transcript file not found.", parent=self.root)
    
    def _export_history_item(self):
        """Export selected history item."""
        self._view_history_transcript()  # Reuses export logic
    
    def _remove_history_item(self):
        """Remove item from history."""
        selection = self.history_tree.selection()
        if not selection:
            return
        
        item = selection[0]
        values = self.history_tree.item(item, "values")
        filename = values[1]
        
        # Remove from history_items
        self.history_items = [h for h in self.history_items if h["file"] != filename]
        self._save_history()
        self._populate_history_tree()
    
    def _clear_history(self):
        if messagebox.askyesno("Clear History", "Remove all history entries?", parent=self.root):
            self.history_items = []
            self._save_history()
            self._populate_history_tree()
    
    def _export_history(self):
        """Export entire history as CSV."""
        path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("JSON", "*.json"), ("All", "*.*")],
            title="Export History"
        )
        if not path:
            return
        
        try:
            if path.endswith(".json"):
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(self.history_items, f, ensure_ascii=False, indent=2)
            else:
                import csv
                with open(path, "w", newline="", encoding="utf-8") as f:
                    writer = csv.DictWriter(f, fieldnames=["time", "file", "type", "status", "duration", "language"])
                    writer.writeheader()
                    writer.writerows(self.history_items)
            messagebox.showinfo("Export Complete", f"History exported to {path}", parent=self.root)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e), parent=self.root)
    
    # ─── Settings Dialog ─────────────────────────────────────────────────
    
    def _open_settings(self):
        """Open settings dialog."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Settings")
        dialog.geometry("500x500")
        dialog.transient(self.root)
        dialog.grab_set()
        
        notebook = ttk.Notebook(dialog, padding=10)
        notebook.pack(fill="both", expand=True)
        
        # General tab
        general = ttk.Frame(notebook, padding=10)
        notebook.add(general, text="General")
        
        ttk.Label(general, text="Theme:").grid(row=0, column=0, sticky="w", pady=5)
        theme_var = tk.StringVar(value=self.settings.get("theme"))
        theme_combo = ttk.Combobox(general, textvariable=theme_var, 
                                  values=["darkly", "superhero", "solar", "cyborg", "flatly", "litera", "minty", "lumen"],
                                  state="readonly")
        theme_combo.grid(row=0, column=1, sticky="ew", pady=5, padx=5)
        
        ttk.Label(general, text="Font Size:").grid(row=1, column=0, sticky="w", pady=5)
        font_var = tk.IntVar(value=self.settings.get("font_size", 10))
        font_spin = ttk.Spinbox(general, from_=8, to=18, textvariable=font_var, width=5)
        font_spin.grid(row=1, column=1, sticky="w", pady=5, padx=5)
        
        remember_var = tk.BooleanVar(value=self.settings.get("remember_window_state", True))
        ttk.Checkbutton(general, text="Remember window size & position", variable=remember_var).grid(row=2, column=0, columnspan=2, sticky="w", pady=5)
        
        confirm_var = tk.BooleanVar(value=self.settings.get("confirm_before_cancel", True))
        ttk.Checkbutton(general, text="Confirm before cancelling transcription", variable=confirm_var).grid(row=3, column=0, columnspan=2, sticky="w", pady=5)
        
        # Transcription tab
        trans = ttk.Frame(notebook, padding=10)
        notebook.add(trans, text="Transcription")
        
        ttk.Label(trans, text="Default Language:").grid(row=0, column=0, sticky="w", pady=5)
        lang_var = tk.StringVar(value=self.settings.get("language", t.AUTO_LANGUAGE_LABEL))
        lang_combo = ttk.Combobox(trans, textvariable=lang_var, values=t.LANGUAGE_COMBO_VALUES, state="readonly")
        lang_combo.grid(row=0, column=1, sticky="ew", pady=5, padx=5)
        
        auto_save_var = tk.BooleanVar(value=self.settings.get("auto_save_speaker", True))
        ttk.Checkbutton(trans, text="Auto-save speaker transcripts", variable=auto_save_var).grid(row=1, column=0, columnspan=2, sticky="w", pady=5)
        
        ttk.Label(trans, text="Default Export Format:").grid(row=2, column=0, sticky="w", pady=5)
        fmt_var = tk.StringVar(value=self.settings.get("default_export_format", "txt"))
        fmt_combo = ttk.Combobox(trans, textvariable=fmt_var, values=[f[1] for f in EXPORT_FORMATS], state="readonly")
        fmt_combo.grid(row=2, column=1, sticky="ew", pady=5, padx=5)
        
        # Audio tab
        audio = ttk.Frame(notebook, padding=10)
        notebook.add(audio, text="Audio")
        
        preview_var = tk.BooleanVar(value=self.settings.get("audio_preview_enabled", PYGAME_AVAILABLE))
        preview_check = ttk.Checkbutton(audio, text="Enable audio preview (requires pygame)", variable=preview_var)
        preview_check.grid(row=0, column=0, columnspan=2, sticky="w", pady=5)
        if not PYGAME_AVAILABLE:
            preview_check.configure(state="disabled")
        
        # Buttons
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill="x", padx=10, pady=10)
        
        def save_settings():
            self.settings["theme"] = theme_var.get()
            self.settings["font_size"] = font_var.get()
            self.settings["remember_window_state"] = remember_var.get()
            self.settings["confirm_before_cancel"] = confirm_var.get()
            self.settings["language"] = lang_var.get()
            self.settings["auto_save_speaker"] = auto_save_var.get()
            self.settings["default_export_format"] = fmt_var.get()
            self.settings["audio_preview_enabled"] = preview_var.get()
            
            # Apply theme immediately
            if TTKBOOTSTRAP_AVAILABLE:
                try:
                    theme_to_use = theme_var.get()
                    if theme_to_use in self.root.style.theme_names():
                        self.root.style.theme_use(theme_to_use)
                except Exception:
                    pass
            
            self._apply_settings()
            dialog.destroy()
        
        ttk.Button(btn_frame, text="Save", command=save_settings, bootstyle="success" if TTKBOOTSTRAP_AVAILABLE else None).pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side="right", padx=5)
        ttk.Button(btn_frame, text="Reset to Defaults", command=lambda: self._reset_settings(dialog)).pack(side="left", padx=5)
        
        general.columnconfigure(1, weight=1)
        trans.columnconfigure(1, weight=1)
    
    def _reset_settings(self, dialog):
        """Reset settings to defaults."""
        if messagebox.askyesno("Reset Settings", "Restore all settings to defaults?", parent=dialog):
            self.settings._data = DEFAULT_SETTINGS.copy()
            self.settings.save()
            self._apply_settings()
            dialog.destroy()
            messagebox.showinfo("Settings Reset", "Settings restored. Restart for full effect.", parent=self.root)
    
    # ─── Utility Methods ─────────────────────────────────────────────────
    
    def get_selected_language_code(self) -> str | None:
        label = self.language_var.get().strip()
        if label == t.AUTO_LANGUAGE_LABEL:
            return None
        code = t.LANGUAGE_LABEL_TO_CODE.get(label)
        if code is None:
            raise ValueError("Select a supported language from the list.")
        return code
    
    def _start_elapsed_timer(self):
        self._update_elapsed_time()
        self.root.after(1000, self._start_elapsed_timer)
    
    def _update_elapsed_time(self):
        if self.transcription_start_time is not None:
            elapsed = time.time() - self.transcription_start_time
            minutes = int(elapsed // 60)
            seconds = int(elapsed % 60)
            if minutes > 0:
                self.batch_elapsed_label.configure(text=f"⏱ Elapsed: {minutes}m {seconds}s")
            else:
                self.batch_elapsed_label.configure(text=f"⏱ Elapsed: {seconds}s")
        else:
            self.batch_elapsed_label.configure(text="")
    
    def _poll_queue(self):
        try:
            while True:
                event = self.event_queue.get_nowait()
                self._handle_event(event)
        except queue.Empty:
            pass
        self.root.after(100, self._poll_queue)
    
    def _format_size(self, bytes_: int) -> str:
        for unit in ["B", "KB", "MB", "GB"]:
            if bytes_ < 1024:
                return f"{bytes_:.1f} {unit}"
            bytes_ /= 1024
        return f"{bytes_:.1f} TB"
    
    def _format_duration(self, seconds: float) -> str:
        if seconds < 60:
            return f"{seconds:.0f}s"
        elif seconds < 3600:
            return f"{int(seconds//60)}m {int(seconds%60)}s"
        else:
            h = int(seconds // 3600)
            m = int((seconds % 3600) // 60)
            return f"{h}h {m}m"
    
    def _copy_to_clipboard(self, text: str, parent=None):
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        if parent:
            messagebox.showinfo("Copied", "Text copied to clipboard.", parent=parent)
    
    def _show_search_dialog(self, parent, text_widget):
        """Show search dialog for transcript."""
        search_dialog = tk.Toplevel(parent)
        search_dialog.title("Search")
        search_dialog.geometry("300x100")
        search_dialog.transient(parent)
        
        ttk.Label(search_dialog, text="Find:").pack(side="left", padx=5, pady=10)
        search_var = tk.StringVar()
        entry = ttk.Entry(search_dialog, textvariable=search_var, width=25)
        entry.pack(side="left", padx=5, pady=10)
        entry.focus_set()
        
        def do_search():
            text_widget.tag_remove("search", "1.0", "end")
            query = search_var.get()
            if not query:
                return
            start = "1.0"
            while True:
                pos = text_widget.search(query, start, "end", nocase=True)
                if not pos:
                    break
                end = f"{pos}+{len(query)}c"
                text_widget.tag_add("search", pos, end)
                start = end
            text_widget.tag_configure("search", background="yellow", foreground="black")
            if text_widget.tag_ranges("search"):
                text_widget.see(text_widget.tag_ranges("search")[0])
        
        ttk.Button(search_dialog, text="Find", command=do_search).pack(side="left", padx=5)
        entry.bind("<Return>", lambda e: do_search())
    
    def _open_url(self, url: str):
        import webbrowser
        webbrowser.open(url)


# ─── Entry Point ─────────────────────────────────────────────────────────

def main():
    # Use ttkbootstrap Window for themed UI if available, otherwise fall back
    if TTKBOOTSTRAP_AVAILABLE:
        root = tb.Window(themename="darkly")
    elif DND_AVAILABLE:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    
    root.title("Audio Transcription")
    root.minsize(900, 600)
    
    # Set icon if available
    try:
        root.iconbitmap("icon.ico")
    except Exception:
        pass
    
    t.ensure_directories()
    app = TranscriptionApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()